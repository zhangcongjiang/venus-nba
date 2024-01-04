from bs4 import BeautifulSoup
from docx import Document
from markdown import markdown
import os


def md_to_doc(md_file, docx_file):
    # 读取 Markdown 文件内容
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

        # 将 Markdown 内容转换为 HTML
    html = markdown(md_content)

    # 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html, 'html.parser')

    # 创建一个新的 Word 文档
    doc = Document()

    # 遍历解析后的 HTML 内容，并添加到 Word 文档中
    for element in soup.children:
        if element.name == 'p' and element.next_element.name != 'img':  # 处理段落
            doc.add_paragraph(element.get_text())
        elif element.next_element.name == 'img':  # 处理图片
            image_path = element.next_element['src']  # 获取图片路径
            if os.path.exists(image_path):  # 确保图片文件存在
                doc.add_picture(image_path)
                # 可以根据需要添加对其他元素的处理逻辑

    # 保存 Word 文档
    doc.save(docx_file)
